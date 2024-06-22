from typing import List
from io import BytesIO
import docx
from langchain_core.documents import Document

class DOCXLoader:
    def __init__(self):
        """
        Initializes the DOCXLoader.
        """
        pass

    def load(self, file_paths: List[str]) -> List[Document]:
        """
        Loads DOCX files and converts them into Document objects.

        :param file_paths: List of file paths representing the files to load.
        :return: List of Document objects.
        """
        documents = []
        for file_path in file_paths:
            file_type = file_path.split(".")[-1]

            if file_type.lower() not in ["docx"]:
                raise ValueError(f"Unsupported file type: {file_type}")

            with open(file_path, 'rb') as f:
                bytes_io_obj = BytesIO(f.read())

            doc_content = self.read_docx(bytes_io_obj)

            metadata = {"source": "DOCX", "file_name": file_path}
            doc = Document(page_content=doc_content, metadata=metadata)
            documents.append(doc)

        return documents

    def read_docx(self, bytes_io_obj: BytesIO) -> str:
        """
        Reads the content of a DOCX file from a BytesIO object.

        :param bytes_io_obj: BytesIO object containing DOCX file content.
        :return: Text content of the DOCX file.
        """
        doc = docx.Document(bytes_io_obj)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

# Example usage
if __name__ == "__main__":
    file_paths = ["C:/Users/Fahira/Downloads/student_handout_23sep22.docx"]
    loader = DOCXLoader()
    documents = loader.load(file_paths)
    for doc in documents:
        print(doc.page_content)
