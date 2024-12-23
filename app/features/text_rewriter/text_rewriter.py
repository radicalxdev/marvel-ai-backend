import sys
import os
from dotenv import load_dotenv
load_dotenv(dotenv_path="/Users/timothydrew92/Desktop/Reality Ai Lab/marvel_ai_backend/app/features/text_rewriter/text_rewriter.env")

# Dynamically add the project root to PYTHONPATH
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../")))

# Load environment variables
load_dotenv()

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise Exception("Google API Key is not set. Please check your .env file.")

# Extract Google Cloud API variables
google_api_key = os.getenv("GOOGLE_API_KEY")
project_id = os.getenv("PROJECT_ID")

# Validate environment variables
if not google_api_key or not project_id:
    raise ValueError("Missing required environment variables. Ensure GOOGLE_API_KEY and PROJECT_ID are set in text_rewriter.env.")

from typing import Union
from docx import Document
import csv
import json
from fastapi import FastAPI
from app.features.text_rewriter.router import router as text_rewriter_router  # Import the router


# Create FastAPI instance
app = FastAPI()
app.include_router(text_rewriter_router, prefix="/text-rewriter")

# Example endpoint for testing
@app.get("/")
def read_root():
    return {"message": "Hello, Reality AI Lab!"}

# Helper Functions
def process_text_input(text: str) -> str:
    """Handles direct text input."""
    return text

def process_file_input(file_path: str) -> str:
    """Processes different file types and extracts text content."""
    _, file_extension = os.path.splitext(file_path)
    
    try:
        if file_extension == '.txt':
            with open(file_path, 'r') as file:
                return file.read()
        elif file_extension == '.docx':
            doc = Document(file_path)
            return '\n'.join([para.text for para in doc.paragraphs])
        elif file_extension == '.csv':
            with open(file_path, 'r') as file:
                return '\n'.join([','.join(row) for row in csv.reader(file)])
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    except Exception as e:
        return f"Error processing file: {str(e)}"

def rewrite_text(content: str, instructions: str) -> str:
    """Rewrites text based on provided instructions."""
    # Placeholder for actual AI rewriting logic
    return f"Rewritten content: [{content}] based on instructions: [{instructions}]"

def export_output(content: str, output_format: str, output_path: str) -> str:
    """Exports the rewritten text in the specified format."""
    try:
        if output_format == 'txt':
            with open(output_path, 'w') as file:
                file.write(content)
        elif output_format == 'docx':
            doc = Document()
            doc.add_paragraph(content)
            doc.save(output_path)
        elif output_format == 'pdf':
            # Placeholder: PDF handling can be added with PyPDF2 or similar library
            raise NotImplementedError("PDF export is not yet supported.")
        else:
            raise ValueError("Unsupported output format.")
        return f"File exported to {output_path}"
    except Exception as e:
        return f"Error exporting file: {str(e)}"

# Main Functionality
def text_rewriter(input_data: Union[str, dict], instructions: str, output_format: str = 'txt', output_path: str = './output.txt') -> str:
    """
    Executes the text rewriting functionality.
    
    Args:
        input_data (Union[str, dict]): Original text or file input.
        instructions (str): Rewriting instructions.
        output_format (str): Desired output format ('txt', 'docx', 'pdf').
        output_path (str): Path to save the output file.
    
    Returns:
        str: Status or rewritten content.
    """
    try:
        if isinstance(input_data, str) and os.path.isfile(input_data):
            content = process_file_input(input_data)
        else:
            content = process_text_input(input_data)
        
        rewritten_content = rewrite_text(content, instructions)
        export_status = export_output(rewritten_content, output_format, output_path)
        
        return f"Rewriting complete. {export_status}"
    except Exception as e:
        return f"Error: {str(e)}"

# Start the FastAPI app
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
