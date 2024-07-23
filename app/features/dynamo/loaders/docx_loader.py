import docx
from fastapi import UploadFile
from langchain_core.documents import Document

class DOCXLoader:
    def __init__(self, files: list[UploadFile]):
        self.files = files

    def load(self) -> list:
        documents = []

        for upload_file in self.files:
            doc = docx.Document(upload_file.file)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            content = "\n".join(full_text)
            metadata = {"source": upload_file.filename}
            doc = Document(page_content=content, metadata=metadata)
            documents.append(doc)

        return documents